import logging

from pathlib import Path

from typing import List, Dict, Any

from docx import Document

from docx.text.paragraph import Paragraph

from docx.table import Table

import config

from processors.docx_processor import DOCXProcessor

from replacers.faker_mapper import FakerMapper

from redaction.image_blur import blur_logos_in_pil

logger = logging.getLogger("byepii")

class DOCXRedactor:

    def __init__(self):

        self.mapper = FakerMapper()

    def redact_docx(self, file_path: str, orchestrator: Any) -> Dict[str, Any]:

        docx_path = Path(file_path)

        processor = DOCXProcessor(docx_path)

        doc = processor.document

        all_detected_entities = []

        orig_previews = [p.text for p in doc.paragraphs if p.text.strip()]

        paragraphs_to_process = []

        last_paragraph_text = ""

        skip_titles = [

            "conventional and general terms and abbreviations",

            "technical and industry related terms",

            "key financial and operating metrics used in this red herring prospectus"

        ]

        def should_skip_table_by_title(title_text: str) -> bool:

            if not title_text:

                return False

            normalized = " ".join(title_text.lower().strip().split())

            for skip_t in skip_titles:

                if skip_t in normalized:

                    return True

            return False

        def iter_block_items(parent):

            parent_elm = parent.element.body

            for child in parent_elm.iterchildren():

                if child.tag.endswith('p'):

                    yield Paragraph(child, parent)

                elif child.tag.endswith('tbl'):

                    yield Table(child, parent)

        for item in iter_block_items(doc):

            if isinstance(item, Paragraph):

                txt = item.text.strip()

                if txt:

                    last_paragraph_text = txt

                paragraphs_to_process.append(item)

            elif isinstance(item, Table):

                if should_skip_table_by_title(last_paragraph_text):

                    logger.info(f"Skipping table under title: '{last_paragraph_text}'")

                    continue

                term_cols = set()

                if len(item.rows) > 0:

                    header_row = item.rows[0]

                    for idx, cell in enumerate(header_row.cells):

                        col_hdr = cell.text.strip().lower()

                        if col_hdr in ("term", "abbreviations"):

                            term_cols.add(idx)

                for row in item.rows:

                    for col_idx, cell in enumerate(row.cells):

                        if col_idx in term_cols:

                            continue                                                             

                        for p in cell.paragraphs:

                            paragraphs_to_process.append(p)

        for section in doc.sections:

            if section.header:

                for p in section.header.paragraphs:

                    paragraphs_to_process.append(p)

            if section.footer:

                for p in section.footer.paragraphs:

                    paragraphs_to_process.append(p)

        for p_idx, p in enumerate(paragraphs_to_process):

            text = p.text.strip()

            if not text:

                continue

            text_entities = orchestrator.detect_text(text)

            if not text_entities:

                continue

            page_ents = []

            for ent in text_entities:

                ent["page"] = 1                                                

                ent["replacement"] = self.mapper.replace(ent["text"], ent["label"])

                page_ents.append(ent)

            if page_ents:

                all_detected_entities.extend(page_ents)

                entity_mappings = {}

                for ent in page_ents:

                    entity_mappings[ent["text"]] = ent["replacement"]

                self._replace_text_in_paragraph(p, entity_mappings)

        if getattr(config, "REDACT_IMAGES", True):
            try:
                self._blur_inline_images(doc)
            except Exception as e:
                logger.warning(f"DOCX inline image blurring failed: {e}")

        output_docx_name = "output.docx"
        output_docx_path = config.OUTPUT_DIR / output_docx_name

        doc.save(str(output_docx_path))

        redacted_previews = [p.text for p in doc.paragraphs if p.text.strip()]

        return {

            "entities": all_detected_entities,

            "original_pages": orig_previews if orig_previews else ["No text content."],

            "redacted_pages": redacted_previews if redacted_previews else ["No text content."],

            "redacted_file_path": str(output_docx_path)

        }

    def _blur_inline_images(self, doc) -> int:

        from PIL import Image

        import io

        total = 0

        rels = doc.part.rels

        for rel_id, rel in list(rels.items()):

            if "image" not in rel.reltype:

                continue

            try:

                img_bytes = rel.target_part.blob

                pil_img = Image.open(io.BytesIO(img_bytes)).convert("RGB")

                blurred = blur_logos_in_pil(pil_img)

                if blurred is not pil_img:

                    buf = io.BytesIO()

                    blurred.save(buf, format="PNG")

                    rel.target_part._blob = buf.getvalue()

                    total += 1

                    logger.info(f"Blurred inline DOCX image (rel: {rel_id})")

            except Exception as e:

                logger.debug(f"Could not process DOCX image rel {rel_id}: {e}")

        return total

    def _replace_text_in_paragraph(self, p, entity_mappings: Dict[str, str]):

        from docx.enum.text import WD_COLOR_INDEX

        text = p.text

        matches = []

        for orig, fake in entity_mappings.items():

            start_idx = 0

            while True:

                idx = text.find(orig, start_idx)

                if idx == -1:

                    break

                matches.append((idx, idx + len(orig), fake))

                start_idx = idx + len(orig)

        if not matches:

            return

        matches = sorted(matches, key=lambda x: x[0])

        non_overlapping = []

        last_end = -1

        for m in matches:

            if m[0] >= last_end:

                non_overlapping.append(m)

                last_end = m[1]

        p.text = ""

        last_idx = 0

        for start, end, fake in non_overlapping:

            if start > last_idx:

                p.add_run(text[last_idx:start])

            r = p.add_run(fake)

            r.bold = True

            r.font.highlight_color = WD_COLOR_INDEX.YELLOW

            r.font.name = "Courier New"

            last_idx = end

        if last_idx < len(text):

            p.add_run(text[last_idx:])
