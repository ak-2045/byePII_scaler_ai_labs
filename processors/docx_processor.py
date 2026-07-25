from pathlib import Path

from typing import Dict, List

from docx import Document

from PIL import Image

import io

class DOCXProcessor:

    def __init__(self, docx_path: str | Path):

        self.docx_path = str(docx_path)

        self.document = Document(self.docx_path)

    def extract_text(self) -> List[Dict]:

        paragraphs = []

        for index, paragraph in enumerate(self.document.paragraphs):

            text = paragraph.text.strip()

            if text:

                paragraphs.append(

                    {

                        "paragraph": index + 1,

                        "text": text

                    }

                )

        return paragraphs

    def extract_tables(self) -> List[Dict]:

        tables = []

        for table_index, table in enumerate(self.document.tables):

            rows = []

            for row in table.rows:

                rows.append(

                    [cell.text.strip() for cell in row.cells]

                )

            tables.append(

                {

                    "table": table_index + 1,

                    "rows": rows

                }

            )

        return tables

    def extract_images(self) -> List[Dict]:

        images = []

        relationships = self.document.part._rels

        image_index = 0

        for relationship in relationships.values():

            if "image" not in relationship.target_ref:

                continue

            image = Image.open(

                io.BytesIO(

                    relationship.target_part.blob

                )

            )

            images.append(

                {

                    "index": image_index,

                    "image": image,

                    "format": image.format

                }

            )

            image_index += 1

        return images

    def statistics(self) -> Dict:

        word_count = sum(

            len(paragraph.text.split())

            for paragraph in self.document.paragraphs

        )

        return {

            "paragraphs": len(self.document.paragraphs),

            "tables": len(self.document.tables),

            "images": len(self.extract_images()),

            "words": word_count

        }
